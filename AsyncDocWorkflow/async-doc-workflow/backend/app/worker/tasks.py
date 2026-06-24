"""
Background document processing tasks.

Flow:
  1. FastAPI receives upload → saves file → creates DB row (status=uploaded)
  2. FastAPI enqueues process_document.delay(doc_id)  ← returns immediately
  3. Celery worker picks up the task (outside request cycle)
  4. Worker updates DB + publishes Redis Pub/Sub events at each stage
  5. WebSocket endpoint relays events to the browser in real time
"""

import json
import logging
import time
from pathlib import Path

import psycopg2
from celery import Task

from app.config import settings
from app.pubsub.redis_pubsub import publish_progress_sync
from app.worker.celery_app import celery_app

logger = logging.getLogger(__name__)


# ── Sync DB helpers (Celery runs sync) ───────────────────────────

def _build_dsn() -> tuple[str, dict]:
    """
    Build a psycopg2-compatible DSN from settings.DATABASE_URL.
    Strips +asyncpg dialect and moves SSL params to connect kwargs.
    """
    url = settings.DATABASE_URL

    # Remove asyncpg dialect prefix
    url = url.replace("postgresql+asyncpg://", "postgresql://")

    # Split off query string
    extra_kwargs = {}
    if "?" in url:
        base, qs = url.split("?", 1)
        # Parse each param
        for param in qs.split("&"):
            if not param:
                continue
            if param.startswith("sslmode="):
                extra_kwargs["sslmode"] = param.split("=", 1)[1]
            elif param.startswith("channel_binding="):
                pass  # psycopg2 doesn't support this, skip it
        url = base

    return url, extra_kwargs


def _get_sync_conn():
    """Return a plain psycopg2 connection compatible with Neon."""
    dsn, kwargs = _build_dsn()
    logger.info(f"[DB] Connecting with DSN: {dsn[:60]}... kwargs={kwargs}")
    return psycopg2.connect(dsn, **kwargs)


def _update_doc(doc_id: str, **fields) -> None:
    set_clause = ", ".join(f"{k} = %s" for k in fields)
    values = list(fields.values()) + [doc_id]
    try:
        conn = _get_sync_conn()
        with conn.cursor() as cur:
            cur.execute(
                f"UPDATE documents SET {set_clause}, updated_at = NOW() WHERE id = %s",
                values,
            )
        conn.commit()
        conn.close()
        logger.info(f"[DB] Updated doc {doc_id}: {list(fields.keys())}")
    except Exception as e:
        logger.error(f"[DB] Failed to update doc {doc_id}: {e}")
        raise


def _emit(doc_id: str, status: str, progress: int, message: str = "", extracted_data=None):
    """Update DB and publish a progress event."""
    update: dict = {"status": status, "progress": progress}
    if extracted_data is not None:
        update["extracted_data"] = json.dumps(extracted_data)
    if status == "failed":
        update["error_message"] = message

    _update_doc(doc_id, **update)

    publish_progress_sync(doc_id, {
        "doc_id": doc_id,
        "status": status,
        "progress": progress,
        "message": message,
        "extracted_data": extracted_data,
    })
    logger.info(f"[EMIT] doc={doc_id} status={status} progress={progress} msg={message}")


# ── Extraction helpers ────────────────────────────────────────────

def _extract_pdf(path: Path) -> dict:
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(str(path))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        return {
            "page_count": len(reader.pages),
            "raw_text": text[:5000],
            "word_count": len(text.split()),
        }
    except Exception as e:
        logger.error(f"[PDF] extraction error: {e}")
        return {"error": str(e)}


def _extract_docx(path: Path) -> dict:
    try:
        from docx import Document
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        return {
            "paragraph_count": len(doc.paragraphs),
            "raw_text": text[:5000],
            "word_count": len(text.split()),
        }
    except Exception as e:
        logger.error(f"[DOCX] extraction error: {e}")
        return {"error": str(e)}


def _extract_image(path: Path) -> dict:
    try:
        from PIL import Image
        img = Image.open(str(path))
        return {
            "width": img.width,
            "height": img.height,
            "mode": img.mode,
            "format": img.format,
        }
    except Exception as e:
        logger.error(f"[IMAGE] extraction error: {e}")
        return {"error": str(e)}


def _extract_text(path: Path) -> dict:
    try:
        text = path.read_text(errors="replace")
        return {
            "raw_text": text[:5000],
            "line_count": text.count("\n"),
            "word_count": len(text.split()),
        }
    except Exception as e:
        logger.error(f"[TEXT] extraction error: {e}")
        return {"error": str(e)}


EXTRACTORS = {
    "application/pdf": _extract_pdf,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": _extract_docx,
    "image/png": _extract_image,
    "image/jpeg": _extract_image,
    "text/plain": _extract_text,
}


# ── Main Celery task ──────────────────────────────────────────────

@celery_app.task(bind=True, max_retries=3, default_retry_delay=5)
def process_document(self: Task, doc_id: str, file_path: str, mime_type: str) -> dict:
    """
    Process a document outside the request-response cycle.

    Stages:
      10% → validating
      30% → extracting
      70% → post-processing
      90% → saving
     100% → completed
    """
    logger.info(f"[TASK] Starting process_document doc_id={doc_id} mime={mime_type}")
    try:
        path = Path(file_path)

        # Stage 1 — validate
        _emit(doc_id, "processing", 10, "Validating file...")
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        time.sleep(0.5)

        # Stage 2 — extract
        _emit(doc_id, "processing", 30, "Extracting content...")
        extractor = EXTRACTORS.get(mime_type, _extract_text)
        extracted = extractor(path)
        extracted["filename"] = path.name
        extracted["mime_type"] = mime_type
        extracted["file_size_bytes"] = path.stat().st_size
        time.sleep(1)

        # Stage 3 — post-process
        _emit(doc_id, "processing", 70, "Post-processing...")
        time.sleep(0.5)

        # Stage 4 — save
        _emit(doc_id, "processing", 90, "Saving results...", extracted_data=extracted)
        time.sleep(0.3)

        # Stage 5 — done
        _emit(doc_id, "completed", 100, "Done", extracted_data=extracted)

        logger.info(f"[TASK] Completed doc_id={doc_id}")
        return {"doc_id": doc_id, "status": "completed"}

    except Exception as exc:
        logger.error(f"[TASK] Failed doc_id={doc_id} error={exc}")
        try:
            _emit(doc_id, "failed", 0, str(exc))
        except Exception:
            pass
        raise self.retry(exc=exc)